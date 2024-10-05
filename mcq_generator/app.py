import streamlit as st
from pdfminer.high_level import extract_text
from langchain.schema import HumanMessage
from langchain.prompts import PromptTemplate, ChatPromptTemplate, HumanMessagePromptTemplate
from langchain_openai import ChatOpenAI
from langchain.output_parsers import StructuredOutputParser, ResponseSchema
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain_core.prompts import PromptTemplate
from langchain.chains.combine_documents.stuff import StuffDocumentsChain
from langchain.chains.llm import LLMChain
import string
import re
import json
import docx
import pandas as pd
import io

def postprocess_text(text):
    printable = (
        string.ascii_letters +
        string.digits +
        string.punctuation +
        string.whitespace +
        'АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ' +
        'абвгдеёжзийклмнопрстуфхцчшщъыьэюя'
    )
    cleaned_text = ''.join(char for char in text if char in printable)
    cleaned_text = re.sub('\xad(\x0c)*', '', cleaned_text)
    cleaned_text = re.sub('[\n-\x0c]', '', cleaned_text)
    return cleaned_text

def result_to_csv(text):
    cleaned_prediction = re.sub(r'```json|```', '', text).strip()
    json_objects = re.findall(r'\{.*?\}', cleaned_prediction, re.DOTALL)
    parsed_data = [json.loads(obj) for obj in json_objects]
    df = pd.DataFrame(parsed_data)
    # Сохранение в буфер
    output = io.BytesIO()
    df.to_excel(output, index=False)
    output.seek(0)  # Перемещение указателя в начало буфера
    return output

st.title("Упроститель текста и генератор вопросов")
uploaded_file = st.file_uploader("Загрузите DOCX файл", type=["docx"])

if uploaded_file is not None:
    def extract_text_from_docx(file):
        doc = docx.Document(file)
        return '\n'.join([para.text for para in doc.paragraphs])

    full_text = extract_text_from_docx(uploaded_file)

    chunk_size = st.number_input("Введите размер чанка", min_value=100, max_value=10000, value=5000, step=100)
    chunk_overlap = int(chunk_size * 0.2)

    if st.button("Разделить"):
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        st.session_state.chunks = text_splitter.split_text(full_text)

    if 'chunks' in st.session_state:
        chunk_index = st.selectbox("Выберите чанк", range(len(st.session_state.chunks)))
        selected_chunk = st.session_state.chunks[chunk_index]
        st.write("Выбранный чанк:")
        st.text(selected_chunk)

        chat_model = ChatOpenAI(temperature=0, model_name='gpt-4o', api_key=st.secrets["OPENAI_API_KEY"])

        if st.button("Упростить чанк"):
            st.session_state.simplify_chunk = True
            st.session_state.generate_questions = False

        if st.button("Сгенерировать вопросы по чанку"):
            st.session_state.generate_questions = True
            st.session_state.simplify_chunk = False

        if st.session_state.get('simplify_chunk', False):
            doc = [Document(page_content=selected_chunk)]
            prompt_template = """Перепиши текст, упростив его, при этом не используй нумерацию абзацев:
            "{text}".
            """
            prompt = PromptTemplate.from_template(prompt_template)
            llm_chain = LLMChain(llm=chat_model, prompt=prompt)
            stuff_chain = StuffDocumentsChain(llm_chain=llm_chain, document_variable_name="text")
            selected_chunk_summary = stuff_chain.run(doc)
            st.markdown(selected_chunk_summary)

        if st.session_state.get('generate_questions', False):
            processed_text = selected_chunk
            response_schemas = [
                ResponseSchema(name="question", description="Вопрос с множественным выбором ответов, созданный на основе фрагмента входного текста."),
                ResponseSchema(name="option_1", description="Первый вариант ответа на вопрос с множественным выбором. Используйте этот формат: 'a) вариант ответа'"),
                ResponseSchema(name="option_2", description="Второй вариант ответа на вопрос с множественным выбором. Используйте этот формат: 'b) вариант ответа''"),
                ResponseSchema(name="option_3", description="Третий вариант ответа на вопрос с множественным выбором. Используйте этот формат: 'c) вариант ответа''"),
                ResponseSchema(name="answer", description="Правильный ответ на вопрос. Используйте этот формат: 'd) вариант ответа'' или 'b) вариант ответа'', и так далее."),
                ResponseSchema(name="topic_number", description="Номер пункта исходного документа."),
                ResponseSchema(name="topic", description="Текст пункта исходного документа с номером topic_number"),
            ]
            output_parser = StructuredOutputParser.from_response_schemas(response_schemas)
            format_instructions = output_parser.get_format_instructions()
            number_of_questions = int(chunk_size / 1000)
            prompt = ChatPromptTemplate(
                messages=[
                    HumanMessagePromptTemplate.from_template("""Получив текст нормативного документа, сгенерируй из него до {number_of_questions} вопросов с несколькими вариантами ответов с правильным ответом.
                    \n{format_instructions}\n{user_prompt}""")
                ],
                input_variables=["user_prompt"],
                partial_variables={"number_of_questions": number_of_questions, "format_instructions": format_instructions}
            )
            user_query = prompt.format_prompt(user_prompt=processed_text)
            user_query_output = chat_model(user_query.to_messages())
            st.write("Сгенерированные вопросы:")
            result = user_query_output.content
            st.write(result)

            # Получение буфера с данными и добавление кнопки для скачивания
            output_buffer = result_to_csv(result)
            st.download_button(
                label="Скачать файл output.xlsx",
                data=output_buffer,
                file_name="output.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )